#!/usr/bin/env python3
"""Add the approved Reviewer 1 Comment 8 Table S5 to the standalone supplement."""

from __future__ import annotations

import argparse
import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


CURRENT_CONTENTS = (
    "These supplementary tables and figure report ordinal-model robustness analyses "
    "for life satisfaction, sample-alignment diagnostics, and exploratory "
    "analytical-place pathway heterogeneity. Table S1 presents the prespecified "
    "four-category analysis, Table S2 reports sensitivity on the original 0–10 "
    "scale, Table S3 documents sample construction, variable-level missingness, and "
    "exact model denominators, and Table S4 reports place-specific direct and "
    "indirect associations, global heterogeneity tests, and survey-weighted "
    "sensitivity diagnostics. Figure S1 displays the place-specific "
    "indirect-association estimates."
)
REVISED_CONTENTS = (
    "These supplementary tables and figure report ordinal-model robustness analyses "
    "for life satisfaction, sample-alignment diagnostics, unadjusted analytical-place "
    "rural-urban differences, and exploratory analytical-place pathway heterogeneity. "
    "Table S1 presents the prespecified four-category analysis, Table S2 reports "
    "sensitivity on the original 0–10 scale, Table S3 documents sample construction, "
    "variable-level missingness, and exact model denominators, and Table S4 reports "
    "place-specific direct and indirect associations, global heterogeneity tests, and "
    "survey-weighted sensitivity diagnostics. Table S5 reports unadjusted "
    "survey-weighted rural and urban life-satisfaction means and rural-minus-urban "
    "differences across all 23 analytical places. Figure S1 displays the "
    "place-specific indirect-association estimates."
)
TABLE_TITLE = (
    "Table S5. Survey-weighted unadjusted rural-urban differences in life "
    "satisfaction across analytical places"
)
NOTES = (
    "Rural and urban means are survey-weighted estimates on the locked common "
    "complete-case sample (N = 183,685). Difference is rural minus urban. The 95% "
    "confidence interval is HC3 from a survey-weighted OLS regression with rural "
    "residence as the only predictor; no covariates are included. Rows follow the "
    "same UN M49 region and within-region place order as Figure 6. Hong Kong is "
    "treated as a region; “analytical place” covers all 23 units. CI = confidence "
    "interval."
)
WIDTHS = [1050, 1550, 720, 720, 900, 900, 3520]
HEADERS = [
    "UN M49 region",
    "Analytical place",
    "Rural n",
    "Urban n",
    "Rural mean",
    "Urban mean",
    "Difference, R-U [95% CI]",
]


def element(parent, tag: str, **attrs: str):
    node = OxmlElement(tag)
    for name, value in attrs.items():
        node.set(qn(name), value)
    parent.append(node)
    return node


def replace_children(parent, tag: str) -> None:
    for child in list(parent.findall(qn(tag))):
        parent.remove(child)


def set_table_geometry(table) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblLayout", "w:tblInd"):
        replace_children(properties, tag)
    element(properties, "w:tblW", **{"w:type": "dxa", "w:w": "9360"})
    element(properties, "w:tblLayout", **{"w:type": "fixed"})
    element(properties, "w:tblInd", **{"w:w": "0", "w:type": "dxa"})

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in WIDTHS:
        element(grid, "w:gridCol", **{"w:w": str(width)})


def set_row_properties(row, *, repeat: bool = False) -> None:
    properties = row._tr.get_or_add_trPr()
    replace_children(properties, "w:cantSplit")
    element(properties, "w:cantSplit")
    replace_children(properties, "w:tblHeader")
    if repeat:
        element(properties, "w:tblHeader", **{"w:val": "true"})


def set_cell_properties(cell, width: int, *, top: bool, bottom: bool) -> None:
    properties = cell._tc.get_or_add_tcPr()
    replace_children(properties, "w:tcW")
    element(properties, "w:tcW", **{"w:type": "dxa", "w:w": str(width)})
    replace_children(properties, "w:vAlign")
    element(properties, "w:vAlign", **{"w:val": "center"})
    replace_children(properties, "w:tcMar")
    margins = element(properties, "w:tcMar")
    for edge, value in (("top", 55), ("start", 70), ("bottom", 55), ("end", 70)):
        element(margins, f"w:{edge}", **{"w:w": str(value), "w:type": "dxa"})
    replace_children(properties, "w:tcBorders")
    borders = element(properties, "w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if (edge == "top" and top) or (edge == "bottom" and bottom):
            element(
                borders,
                f"w:{edge}",
                **{"w:val": "single", "w:sz": "8", "w:space": "0", "w:color": "000000"},
            )
        else:
            element(borders, f"w:{edge}", **{"w:val": "nil"})
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, *, size: float, bold: bool = False, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, fonts)
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")


def set_cell_text(cell, text: str, *, align, size: float, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.style = "Table"
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)


def add_page_break_paragraph(document: Document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break()
    break_node = paragraph._p.find(".//" + qn("w:br"))
    break_node.set(qn("w:type"), "page")
    return paragraph


def read_rows(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8-sig") as stream:
        rows = list(csv.DictReader(stream))
    if len(rows) != 23:
        raise ValueError(f"Expected 23 Table S5 rows, found {len(rows)}")
    return rows


def build_table(document: Document, rows: list[dict[str, str]]):
    table = document.add_table(rows=2 + len(rows), cols=7)
    set_table_geometry(table)

    title_cell = table.rows[0].cells[0].merge(table.rows[0].cells[-1])
    set_cell_properties(title_cell, 9360, top=True, bottom=False)
    set_cell_text(
        title_cell,
        TABLE_TITLE,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=10.5,
        bold=True,
    )
    set_row_properties(table.rows[0], repeat=True)

    for column, (header, width) in enumerate(zip(HEADERS, WIDTHS, strict=True)):
        cell = table.rows[1].cells[column]
        set_cell_properties(cell, width, top=False, bottom=True)
        set_cell_text(
            cell,
            header,
            align=WD_ALIGN_PARAGRAPH.LEFT if column < 2 else WD_ALIGN_PARAGRAPH.CENTER,
            size=8.0,
        )
    set_row_properties(table.rows[1], repeat=True)

    for row_index, source in enumerate(rows, start=2):
        values = [
            source["UN M49 region"],
            source["Analytical place"],
            f"{int(source['Rural n']):,}",
            f"{int(source['Urban n']):,}",
            source["Rural weighted mean"],
            source["Urban weighted mean"],
            source["Unadjusted difference, R-U [95% CI]"],
        ]
        for column, (value, width) in enumerate(zip(values, WIDTHS, strict=True)):
            cell = table.rows[row_index].cells[column]
            set_cell_properties(
                cell,
                width,
                top=False,
                bottom=row_index == len(rows) + 1,
            )
            set_cell_text(
                cell,
                value,
                align=WD_ALIGN_PARAGRAPH.LEFT if column < 2 else WD_ALIGN_PARAGRAPH.CENTER,
                size=8.0,
            )
        set_row_properties(table.rows[row_index])
    return table


def add_notes(document: Document):
    paragraph = document.add_paragraph(style="Table")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.keep_with_next = False
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    label = paragraph.add_run("Notes. ")
    set_run_font(label, size=9.0, italic=True)
    text = paragraph.add_run(NOTES)
    set_run_font(text, size=9.0)
    return paragraph


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--table-csv", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    document = Document(args.input)
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == CURRENT_CONTENTS]
    if len(matches) != 1:
        raise ValueError(f"Expected one exact contents paragraph, found {len(matches)}")
    if any(TABLE_TITLE in table.cell(0, 0).text for table in document.tables):
        raise ValueError("Table S5 is already present")
    headings = [paragraph for paragraph in document.paragraphs if paragraph.text == "Supplementary Figure"]
    if len(headings) != 1:
        raise ValueError("Could not uniquely locate the Supplementary Figure heading")
    inference = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Inference. The wild-score bootstrap uses 4,999 draws per place")
    ]
    if len(inference) != 1:
        raise ValueError("Could not uniquely locate the Table S4 inference paragraph")

    replace_paragraph_text(matches[0], REVISED_CONTENTS)
    rows = read_rows(args.table_csv)
    page_break = add_page_break_paragraph(document)
    table = build_table(document, rows)
    notes = add_notes(document)

    anchor = inference[0]._p
    anchor.addnext(page_break._p)
    page_break._p.addnext(table._tbl)
    table._tbl.addnext(notes._p)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)

    verified = Document(args.output)
    if len(verified.tables) != 8:
        raise ValueError(f"Expected 8 tables after update, found {len(verified.tables)}")
    if sum(paragraph.text == REVISED_CONTENTS for paragraph in verified.paragraphs) != 1:
        raise ValueError("Revised contents paragraph verification failed")
    new_table = verified.tables[-1]
    if len(new_table.rows) != 25 or len(new_table.columns) != 7:
        raise ValueError("Table S5 dimensions are incorrect")
    if new_table.cell(0, 0).text != TABLE_TITLE:
        raise ValueError("Table S5 title verification failed")
    if not any(paragraph.text == "Notes. " + NOTES for paragraph in verified.paragraphs):
        raise ValueError("Table S5 notes verification failed")
    print(
        f"updated={args.output} tables={len(verified.tables)} "
        f"table_s5_rows={len(new_table.rows)} table_s5_cols={len(new_table.columns)}"
    )


if __name__ == "__main__":
    main()
