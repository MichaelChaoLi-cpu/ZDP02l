#!/usr/bin/env python3
"""Apply or reverse the Reviewer 2 Comment 7 Supplement terminology bundle."""

from __future__ import annotations

import hashlib
import argparse
import os
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Rev/revision/ZDP02l.supplementary.docx"
FIGURE = ROOT / "reports/comment5_place_path_heterogeneity/figure_place_path_heterogeneity.png"


# Longer phrases precede their components. Counts are visible-text occurrence counts in
# the approved 24-object Supplement scope (37 non-overlapping replacement spans).
REPLACEMENTS = (
    ("Within-Place Income Percentile (income_pctile)",
     "Within-Region and Within-Country Income Percentile (income_pctile)", 1),
    ("Within-place income percentile", "Within-region and within-country income percentile", 3),
    ("Global cross-place heterogeneity tests", "Global cross-region and cross-country heterogeneity tests", 1),
    ("cross-place", "cross-region and cross-country", 3),
    ("Hong Kong is treated as a region; 'analytical place' covers all 23 units.",
     "Hong Kong is treated as a region; 'region or country' covers all 23 units.", 1),
    ("Hong Kong is treated as a region; “analytical place” covers all 23 units.",
     "Hong Kong is treated as a region; “region or country” covers all 23 units.", 1),
    ("23 analytical places", "23 regions and countries", 1),
    ("23 place clusters", "23 region and country clusters", 2),
    ("Analytical Place (COUNTRY)", "Region or Country (COUNTRY)", 1),
    ("Analytical place", "Region or country", 3),
    ("each analytical place", "each region or country", 1),
    ("across analytical places", "across regions and countries", 1),
    ("analytical-place", "region-and-country", 6),
    ("Place-specific", "Region- and country-specific", 2),
    ("place-specific", "region- and country-specific", 4),
    ("place-clustered", "region- and country-clustered", 2),
    ("place estimates", "region and country estimates", 1),
    ("individual places", "individual regions and countries", 1),
    ("draws per place", "draws per region or country", 1),
    ("place order", "region and country order", 1),
)


def iter_unique_paragraphs(document: Document):
    seen: set[int] = set()
    paragraphs = list(document.paragraphs)
    paragraphs.extend(
        paragraph
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    for paragraph in paragraphs:
        key = id(paragraph._p)
        if key not in seen:
            seen.add(key)
            yield paragraph


def visible_text(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in iter_unique_paragraphs(document))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--restore-original",
        action="store_true",
        help="Reverse the regions-and-countries bundle after the human narrowed it to the manuscript title.",
    )
    args = parser.parse_args()
    replacements = (
        tuple((new, old, expected) for old, new, expected in REPLACEMENTS)
        if args.restore_original
        else REPLACEMENTS
    )

    if not DOCX.is_file() or not FIGURE.is_file():
        raise FileNotFoundError("Supplement or regenerated Figure S1 is missing")

    document = Document(DOCX)
    before = visible_text(document)
    expected_total = sum(expected for _, _, expected in replacements)
    if expected_total != 37:
        raise AssertionError(f"Replacement manifest totals {expected_total}, expected 37")

    changed_objects = 0
    changed_occurrences = 0
    audit_text = before
    mismatches = {}
    for old, new, expected in replacements:
        actual = audit_text.count(old)
        if actual not in (0, expected):
            mismatches[old] = (actual, f"0 or {expected}")
        audit_text = audit_text.replace(old, new)
    if mismatches:
        raise AssertionError(f"Precondition count mismatch: {mismatches}")

    if any(old in before for old, _, _ in replacements):
        for paragraph in iter_unique_paragraphs(document):
            paragraph_changed = False
            for run in paragraph.runs:
                for old, new, _ in replacements:
                    count = run.text.count(old)
                    if count:
                        run.text = run.text.replace(old, new)
                        changed_occurrences += count
                        paragraph_changed = True
            if paragraph_changed:
                changed_objects += 1

        if (changed_objects, changed_occurrences) not in ((24, 37), (3, 3)):
            raise AssertionError(
                f"Unexpected partial state: {changed_occurrences} spans in {changed_objects} objects"
            )

    if len(document.inline_shapes) != 1:
        raise AssertionError(f"Expected one inline figure, found {len(document.inline_shapes)}")
    blips = document.inline_shapes[0]._inline.xpath(".//a:blip")
    if len(blips) != 1:
        raise AssertionError(f"Expected one image blip, found {len(blips)}")
    rid = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
    image_part = document.part.rels[rid].target_part
    old_image_sha = hashlib.sha256(image_part.blob).hexdigest()
    image_part._blob = FIGURE.read_bytes()
    new_image_sha = hashlib.sha256(image_part.blob).hexdigest()
    expected_image_sha = hashlib.sha256(FIGURE.read_bytes()).hexdigest()
    if new_image_sha != expected_image_sha:
        raise AssertionError("Embedded Figure S1 payload does not match regenerated source")

    fd, temp_name = tempfile.mkstemp(prefix=f".{DOCX.stem}.", suffix=".docx", dir=DOCX.parent)
    os.close(fd)
    staged = Path(temp_name)
    try:
        document.save(staged)
        with zipfile.ZipFile(staged) as archive:
            bad = archive.testzip()
            if bad is not None:
                raise AssertionError(f"Invalid DOCX ZIP member: {bad}")
            for name in archive.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    etree.fromstring(archive.read(name))

        check = Document(staged)
        after = visible_text(check)
        for old, _, _ in replacements:
            if old in after:
                raise AssertionError(f"Old visible term remains after edit: {old!r}")
        for protected in ("COUNTRY", "income_pctile", "Hong Kong is treated as a region"):
            if protected not in after:
                raise AssertionError(f"Protected text was lost: {protected!r}")
        os.replace(staged, DOCX)
    finally:
        if staged.exists():
            staged.unlink()

    print(f"updated={DOCX}")
    print(f"mode={'restore-original' if args.restore_original else 'apply'}")
    print(
        "text_objects=already_applied"
        if changed_occurrences == 0
        else f"text_objects={changed_objects} replacement_spans={changed_occurrences}"
    )
    print(f"figure_rid={rid}")
    print(f"figure_old_sha256={old_image_sha}")
    print(f"figure_new_sha256={new_image_sha}")
    print(f"docx_sha256={hashlib.sha256(DOCX.read_bytes()).hexdigest()}")


if __name__ == "__main__":
    main()
